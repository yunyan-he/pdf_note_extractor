import io
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os

def generate_txt_content(notes):
    output = ""
    for note in notes:
        output += f"[Page {note.page}]\n"
        output += f"Quote: {note.original_text}\n"
        if note.note:
            output += f"Note: {note.note}\n"
        output += "-" * 20 + "\n"
    return output

def generate_docx_bytes(notes):
    doc = Document()
    doc.add_heading('PDF Extraction Notes', 0)

    for note in notes:
        p = doc.add_paragraph()
        p.add_run(f"[Page {note.page}] ").bold = True
        
        if note.original_text:
            p.add_run(f"\nQuote: \"{note.original_text}\"").bold = True
        
        if note.note:
            # Add a new line for the note
            run = p.add_run(f"\nNote: {note.note}")
            run.italic = True
        
        doc.add_paragraph("-" * 40)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def generate_pdf_bytes(notes):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 50
    margin = 50
    
    # Try to register a Chinese font
    # Common paths on macOS for Chinese fonts
    font_path = "/System/Library/Fonts/PingFang.ttc"
    font_name = "Helvetica" # Default fallback
    
    # Check for some common Chinese fonts
    possible_fonts = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Hiragino Sans GB.ttc"
    ]
    
    registered_font = False
    for path in possible_fonts:
        if os.path.exists(path):
            try:
                # PingFang.ttc often contains multiple fonts, we need to pick one face ideally
                # But ReportLab TTFont might handle it or we might need to be specific.
                # Actually, simplified extraction: use a known one if possible.
                # For safety in this demo without complex font enumeration:
                # We will try to rely on the fact that we might output bytes. 
                # If we can't load a font, we warn.
                
                # A safer bet for mac is usually:
                pdfmetrics.registerFont(TTFont('Chinese', path))
                font_name = 'Chinese'
                registered_font = True
                break
            except Exception:
                continue
    
    if not registered_font:
        # Fallback to standard font (will not support Chinese)
        # In a real app we should bundle a ttf font file.
        pass

    c.setFont(font_name, 10)

    for note in notes:
        if y < 100:
            c.showPage()
            c.setFont(font_name, 10)
            y = height - 50
            
        text_object = c.beginText(margin, y)
        text_object.setFont(font_name, 12 if registered_font else 10)
        
        lines = []
        lines.append(f"[Page {note.page}]")
        if note.original_text:
            # Simple word wrap logic (very basic for demo)
            # In production, use ReportLab's Paragraph or more robust wrapping
            raw_text = f"Quote: \"{note.original_text}\""
            # Truncate or split nicely would be better, but detailed wrapping is complex without Platypus
            # We'll just dump it for now, maybe split by newlines if they exist
            lines.extend(raw_text.split('\n'))
            
        if note.note:
            raw_note = f"Note: {note.note}"
            lines.extend(raw_note.split('\n'))
            
        lines.append("-" * 40)
        
        for line in lines:
            if y < 50:
                 c.drawText(text_object)
                 c.showPage()
                 c.setFont(font_name, 12 if registered_font else 10)
                 y = height - 50
                 text_object = c.beginText(margin, y)
            
            # Very basic wrap handling by char count (rough API)
            max_chars = 60
            while len(line) > max_chars:
                chunk = line[:max_chars]
                text_object.textLine(chunk)
                line = line[max_chars:]
                y -= 15
            text_object.textLine(line)
            y -= 15

        c.drawText(text_object)
        y -= 20 # Extra spacing between items

    c.save()
    buffer.seek(0)
    return buffer
